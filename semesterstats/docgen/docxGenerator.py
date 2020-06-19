from ..databaseFunctions import getSummary
from ..Models import DepartmentCodeDictionary
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE

## Legacy -- Do Not Use Unless Dire.


def docxGenerator(Batch: int, Semester: int, Department: str):
    Data = getSummary(Batch, Semester, Department)

    document = Document()

    section = document.sections[-1]
    new_width, new_height = Cm(29.7), Cm(21)
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # Add A section header.

    header = section.header
    headerP = header.paragraphs[0]
    headerP.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    headerRun = headerP.add_run()
    headerRun.add_picture("logo.jpeg", width=Pt(80), height=Pt(40))

    p = document.add_paragraph()

    x = p.add_run("C.M.R.INSTITUTE OF TECHNOLOGY, BANGALORE\n")
    x.bold = True
    x.underline = True
    x.font.size = Pt(16)
    x.font.name = "Times New Roman"
    x = p.add_run(f"Result Analysis of {Batch} batch in Semester {Semester}\n")
    x.bold = True
    x.font.size = Pt(16)
    x.font.name = "Times New Roman"
    x = p.add_run(f"Department: {DepartmentCodeDictionary[Department]}")
    x.font.size = Pt(16)
    x.bold = True
    x.font.name = "Times New Roman"
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    x = p.add_run("Total Appeared:{}  ".format(Data["TotalAttendees"]))
    x.font.size = Pt(11)
    x.font.name = "Times New Roman"
    x.bold = True
    x = p.add_run("No of FCD:{}  ".format(Data["FCD"]))
    x.font.size = Pt(11)
    x.font.name = "Times New Roman"
    x.bold = True
    x = p.add_run("No of FC:{}  ".format(Data["FC"]))
    x.font.size = Pt(11)
    x.font.name = "Times New Roman"
    x.bold = True
    x = p.add_run("No of SC:{}  ".format(Data["SC"]))
    x.font.size = Pt(11)
    x.font.name = "Times New Roman"
    x.bold = True
    x = p.add_run("No of Pass:{}  ".format(Data["Pass"]))
    x.font.size = Pt(11)
    x.font.name = "Times New Roman"
    x.bold = True
    x = p.add_run("No of Failures:{}  ".format(Data["Failures"]))
    x.font.size = Pt(11)
    x.font.name = "Times New Roman"
    x.bold = True
    x = p.add_run("Pass percentage:{}".format(Data["PassPercentage"]))
    x.font.size = Pt(11)
    x.font.name = "Times New Roman"
    x.bold = True

    # Add the Table

    Table = document.add_table(rows=len(Data["EachSubjectDetail"]) + 1, cols=9)
    Table.autofit = False
    # Get Column 1

    Column1 = Table.column_cells(0)
    x = Column1[0]

    for y in Column1[1:]:
        x = x.merge(y)

    # Align Said Cell:
    x.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Add Something to Column 1
    p = x.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    z = p.add_run("Year\n")
    z.underline = True
    z.font.name = "Times New Roman"
    z.font.size = Pt(12)
    z.bold = True
    z = p.add_run("{} {}".format(Batch, Semester))
    z.bold = True
    z.font.name = "Times New Roman"
    z.font.size = Pt(12)

    rows = Table.row_cells(0)
    initialRowHeaderDetail = [
        "Subject Code",
        "Faculty Name",
        "Total no of Attendees",
        "Failures",
        "FCD",
        "FC",
        "SC",
        "Pass%",
    ]

    ColumnSpacingDetail = [
        2,
        2.5,
        8,
        2.1,
        2,
        1.5,
        1.5,
        1.5,
        2,
    ]

    for idx, col in enumerate(Table.columns):
        col.width = Cm(ColumnSpacingDetail[idx])

    for idx, cell in enumerate(rows[1:]):
        p = cell.add_paragraph()
        z = p.add_run(initialRowHeaderDetail[idx])
        z.bold = True
        z.font.size = Pt(11)
        z.font.name = "Times New Roman"
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.width = Pt(len(initialRowHeaderDetail[idx]))

    for rowidx, row in enumerate(Table.rows[1:]):
        row.height = Pt(25)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        cells = row.cells
        dd = Data["EachSubjectDetail"][rowidx]
        values = [
            dd["SubjectCode"],
            dd["FacultyName"] + " PlaceHolder for a long string",
            dd["Attendees"],
            dd["Failures"],
            dd["FCD"],
            dd["FC"],
            dd["SC"],
            dd["PassPercentage"],
        ]
        for idx, cell in enumerate(cells[1:]):
            cell.text = str(values[idx])
            p = cell.paragraphs[0]
            p.style.font.name = "Times New Roman"

            if int(idx) == 1:
                p.style.font.size = Pt(5)
            else:
                p.style.font.size = Pt(12)

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.width = Pt(len(str(values[idx])))

    p = document.add_paragraph("\n\nHOD CSE\t PRINCIPAL")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # p.style.font.bold = True
    p.paragraph_format.tab_stops.add_tab_stop(Cm(22), alignment=WD_TAB_ALIGNMENT.RIGHT)
    document.save("demo.docx")
